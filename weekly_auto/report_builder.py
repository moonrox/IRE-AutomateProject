"""weekly_auto.report_builder - render the weekly status report.

Auto-fills the Progress section from collected change items (commits, ADRs,
files). Blockers/Risks and Next Week are narrative — read from weekly_notes.md
(sections tagged [Blockers] and [Next Week]) so a human can steer them.

Produces both a Markdown string and a .docx file in the standard format:
    Progress: / Blockers/Risks: / Next Week:
"""
from __future__ import annotations

import re
from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: F401 (kept for future use)
from docx.shared import Pt, RGBColor

from .collectors import SourceResult
from .util import WorkWeek, now_stamp

_MAX_COMMITS_PER_SOURCE = 12


# ── narrative notes ──────────────────────────────────────────────────────────

def read_notes(notes_path: Path) -> dict[str, list[str]]:
    """Parse weekly_notes.md into {'blockers': [...], 'next_week': [...]}."""
    out = {"blockers": [], "next_week": []}
    if not notes_path.exists():
        return out
    current = None
    for line in notes_path.read_text(encoding="utf-8", errors="ignore").splitlines():
        low = line.strip().lower()
        if low.startswith("[blockers"):
            current = "blockers"
            continue
        if low.startswith("[next week"):
            current = "next_week"
            continue
        if low.startswith("["):
            current = None
            continue
        item = line.strip().lstrip("-*").strip()
        if current and item:
            out[current].append(item)
    return out


# ── aggregation ──────────────────────────────────────────────────────────────

def _group(results: list[SourceResult]) -> tuple[dict, list, list, list, list]:
    """Return (commits_by_source, adr_items, email_items, meeting_items, warnings)."""
    commits: dict[str, list] = defaultdict(list)
    adrs: list = []
    emails: list = []
    meetings: list = []
    warnings: list = []
    for r in results:
        if r.warning:
            warnings.append(f"{r.name}: {r.warning}")
        for it in r.items:
            if it.category == "commit":
                commits[r.name].append(it)
            elif it.category == "adr":
                adrs.append(it)
            elif it.category == "email":
                emails.append(it)
            elif it.category == "meeting":
                meetings.append(it)
            elif it.category == "file":
                commits[r.name].append(it)
    return commits, adrs, emails, meetings, warnings


_MAX_EMAILS = 15
_MAX_MEETINGS = 20


def _progress_lines(results: list[SourceResult]) -> list[str]:
    commits, adrs, emails, meetings, _ = _group(results)
    lines: list[str] = []
    if adrs:
        lines.append("New Architecture Decision Records (ADRs) added this week:")
        for a in sorted(adrs, key=lambda x: x.ref):
            lines.append(f"    - {a.title} ({a.date})")
    for source, items in commits.items():
        if not items:
            continue
        shown = items[:_MAX_COMMITS_PER_SOURCE]
        extra = len(items) - len(shown)
        header = f"{source} - {len(items)} change(s) this week:"
        lines.append(header)
        for it in shown:
            lines.append(f"    - {it.title} ({it.date})")
        if extra > 0:
            lines.append(f"    - ...and {extra} more")
    if meetings:
        shown = sorted(meetings, key=lambda x: x.date)[:_MAX_MEETINGS]
        extra = len(meetings) - len(shown)
        lines.append(f"Meetings this week ({len(meetings)}):")
        for it in shown:
            lines.append(f"    - {it.title} ({it.date})")
        if extra > 0:
            lines.append(f"    - ...and {extra} more")
    if emails:
        shown = emails[:_MAX_EMAILS]
        extra = len(emails) - len(shown)
        lines.append(f"Relevant email activity this week ({len(emails)}):")
        for it in shown:
            lines.append(f"    - {it.title} ({it.date})")
        if extra > 0:
            lines.append(f"    - ...and {extra} more")
    if not lines:
        lines.append("No source activity detected in the work-week window.")
    return lines


# ── markdown ─────────────────────────────────────────────────────────────────

def build_markdown(ww: WorkWeek, author: str, results: list[SourceResult],
                   notes: dict[str, list[str]]) -> str:
    _, _, _, _, warnings = _group(results)
    out: list[str] = []
    out.append(f"# {ww.human}")
    out.append(f"{author} | Infrastructure Reliability Engineering")
    out.append(f"_Generated {now_stamp()} by weekly_auto_")
    out.append("")
    out.append("## Progress")
    out.append("")
    for line in _progress_lines(results):
        indent = "  " if line.startswith("    ") else ""
        out.append(f"{indent}- {line.strip()}" if not line.startswith("    ") else f"  {line.strip()}")
    out.append("")
    out.append("## Blockers / Risks")
    out.append("")
    blockers = notes.get("blockers") or ["None."]
    for b in blockers:
        out.append(f"- {b}")
    out.append("")
    out.append("## Next Week")
    out.append("")
    nxt = notes.get("next_week") or ["TBD."]
    for n in nxt:
        out.append(f"- {n}")
    if warnings:
        out.append("")
        out.append("## Source Coverage Notes")
        out.append("")
        for w in warnings:
            out.append(f"- {w}")
    out.append("")
    return "\n".join(out)


# ── docx ─────────────────────────────────────────────────────────────────────

def build_docx(ww: WorkWeek, author: str, author_title: str,
               results: list[SourceResult], notes: dict[str, list[str]],
               out_path: Path) -> Path:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    t = doc.add_paragraph()
    r = t.add_run(f"{ww.label} Weekly Status - {author}")
    r.bold = True
    r.font.size = Pt(16)

    sub = doc.add_paragraph()
    sr = sub.add_run(author_title)
    sr.italic = True
    sr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    meta = doc.add_paragraph()
    meta.add_run(f"Work Week: {ww.human}").bold = True
    doc.add_paragraph()

    def section(heading: str) -> None:
        h = doc.add_paragraph()
        hr = h.add_run(heading)
        hr.bold = True
        hr.font.size = Pt(12)

    section("Progress:")
    for line in _progress_lines(results):
        if line.startswith("    "):
            p = doc.add_paragraph(line.strip(), style="List Bullet 2")
        else:
            p = doc.add_paragraph(line.strip(), style="List Bullet")
    doc.add_paragraph()

    section("Blockers/Risks:")
    for b in (notes.get("blockers") or ["None."]):
        doc.add_paragraph(b, style="List Bullet")
    doc.add_paragraph()

    section("Next Week:")
    for n in (notes.get("next_week") or ["TBD."]):
        doc.add_paragraph(n, style="List Bullet")

    _, _, _, _, warnings = _group(results)
    if warnings:
        doc.add_paragraph()
        section("Source Coverage Notes:")
        for w in warnings:
            doc.add_paragraph(w, style="List Bullet")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


def markdown_to_html(md: str) -> str:
    """Very small Markdown-to-HTML shim for the email body (headings + bullets)."""
    html: list[str] = ["<html><body style=\"font-family:Calibri,Arial,sans-serif\">"]
    in_list = False
    for line in md.splitlines():
        if line.startswith("# "):
            if in_list:
                html.append("</ul>"); in_list = False
            html.append(f"<h2>{_esc(line[2:])}</h2>")
        elif line.startswith("## "):
            if in_list:
                html.append("</ul>"); in_list = False
            html.append(f"<h3>{_esc(line[3:])}</h3>")
        elif line.strip().startswith("-"):
            if not in_list:
                html.append("<ul>"); in_list = True
            html.append(f"<li>{_esc(line.strip()[1:].strip())}</li>")
        elif line.strip():
            if in_list:
                html.append("</ul>"); in_list = False
            html.append(f"<p>{_esc(line)}</p>")
    if in_list:
        html.append("</ul>")
    html.append("</body></html>")
    return "\n".join(html)


def _esc(s: str) -> str:
    return (s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            .replace("**", ""))
